import io
from typing import Dict, Any
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from datetime import datetime
import logging

def setup_pdf_styles():
    """Setup custom styles for PDF generation."""
    styles = getSampleStyleSheet()
    
    # Title style
    styles.add(ParagraphStyle(
        name='CustomTitle',
        parent=styles['Title'],
        fontSize=24,
        spaceAfter=30,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#2E8B57')
    ))
    
    # Heading styles
    styles.add(ParagraphStyle(
        name='CustomHeading1',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=12,
        spaceBefore=20,
        textColor=colors.HexColor('#2E8B57')
    ))
    
    styles.add(ParagraphStyle(
        name='CustomHeading2',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=10,
        spaceBefore=15,
        textColor=colors.HexColor('#4682B4')
    ))
    
    # Code style
    styles.add(ParagraphStyle(
        name='CodeStyle',
        parent=styles['Code'],
        fontSize=10,
        fontName='Courier',
        backgroundColor=colors.HexColor('#F5F5F5'),
        leftIndent=20,
        rightIndent=20,
        spaceAfter=10
    ))
    
    return styles

def generate_pdf_document(documentation: Dict[str, Any], repo_info: Dict[str, str]) -> io.BytesIO:
    """Generate PDF documentation."""
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        styles = setup_pdf_styles()
        
        # Build document content
        story = []
        
        # Title page
        story.append(Paragraph(f"Architecture Documentation", styles['CustomTitle']))
        story.append(Spacer(1, 20))
        story.append(Paragraph(f"Repository: {repo_info['full_name']}", styles['Heading2']))
        story.append(Paragraph(f"Description: {repo_info['description']}", styles['Normal']))
        story.append(Spacer(1, 10))
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
        story.append(Spacer(1, 30))
        
        # Repository statistics
        if 'key_statistics' in documentation:
            stats = documentation['key_statistics']
            story.append(Paragraph("Repository Statistics", styles['CustomHeading1']))
            
            stats_data = [
                ['Metric', 'Value'],
                ['Total Files', str(stats.get('total_files', 'N/A'))],
                ['Code Files', str(stats.get('code_files', 'N/A'))],
                ['Analyzed Files', str(stats.get('analyzed_files', 'N/A'))],
                ['Skipped Files', str(stats.get('skipped_files', 'N/A'))],
                ['Main Language', stats.get('main_language', 'N/A')],
                ['Modules Count', str(stats.get('modules_count', 'N/A'))],
                ['Stars', str(repo_info.get('stars', 'N/A'))],
                ['Forks', str(repo_info.get('forks', 'N/A'))]
            ]
            
            stats_table = Table(stats_data, colWidths=[2*inch, 2*inch])
            stats_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E8B57')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(stats_table)
            story.append(Spacer(1, 20))
        
        # Overview section
        if documentation.get('overview'):
            story.append(Paragraph("Overview", styles['CustomHeading1']))
            story.append(Paragraph(documentation['overview'], styles['Normal']))
            story.append(Spacer(1, 20))
        
        # Architecture section
        if documentation.get('architecture'):
            story.append(Paragraph("High-Level Architecture", styles['CustomHeading1']))
            story.append(Paragraph(documentation['architecture'], styles['Normal']))
            story.append(Spacer(1, 15))
            
            # Architecture diagram
            if documentation.get('architecture_diagram'):
                story.append(Paragraph("Architecture Diagram", styles['CustomHeading2']))
                story.append(Paragraph("Mermaid Diagram Code:", styles['Normal']))
                story.append(Paragraph(documentation['architecture_diagram'], styles['CodeStyle']))
                story.append(Spacer(1, 15))
        
        # Modules section
        if documentation.get('modules'):
            story.append(Paragraph("Module Breakdown", styles['CustomHeading1']))
            story.append(Paragraph(documentation['modules'], styles['Normal']))
            story.append(Spacer(1, 20))
        
        # Data Flow section
        if documentation.get('data_flow'):
            story.append(Paragraph("Data Flow Analysis", styles['CustomHeading1']))
            story.append(Paragraph(documentation['data_flow'], styles['Normal']))
            story.append(Spacer(1, 15))
            
            # Data flow diagram
            if documentation.get('data_flow_diagram'):
                story.append(Paragraph("Data Flow Diagram", styles['CustomHeading2']))
                story.append(Paragraph("Mermaid Diagram Code:", styles['Normal']))
                story.append(Paragraph(documentation['data_flow_diagram'], styles['CodeStyle']))
                story.append(Spacer(1, 15))
        
        # Design Patterns section
        if documentation.get('design_patterns'):
            story.append(Paragraph("Design Patterns", styles['CustomHeading1']))
            story.append(Paragraph(documentation['design_patterns'], styles['Normal']))
            story.append(Spacer(1, 20))
        
        # Build PDF
        doc.build(story)
        
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        logging.error(f"Error generating PDF: {e}")
        raise

def generate_docx_document(documentation: Dict[str, Any], repo_info: Dict[str, str]) -> io.BytesIO:
    """Generate DOCX documentation."""
    try:
        document = Document()
        
        # Title
        title = document.add_heading('Architecture Documentation', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Repository info
        document.add_heading(f"Repository: {repo_info['full_name']}", level=1)
        document.add_paragraph(f"Description: {repo_info['description']}")
        document.add_paragraph(f"Language: {repo_info['language']}")
        document.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
        
        # Repository statistics
        if 'key_statistics' in documentation:
            document.add_heading('Repository Statistics', level=1)
            stats = documentation['key_statistics']
            
            # Create table for statistics
            table = document.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = 'Value'
            
            # Add data rows
            stats_data = [
                ('Total Files', str(stats.get('total_files', 'N/A'))),
                ('Code Files', str(stats.get('code_files', 'N/A'))),
                ('Analyzed Files', str(stats.get('analyzed_files', 'N/A'))),
                ('Skipped Files', str(stats.get('skipped_files', 'N/A'))),
                ('Main Language', stats.get('main_language', 'N/A')),
                ('Modules Count', str(stats.get('modules_count', 'N/A'))),
                ('Stars', str(repo_info.get('stars', 'N/A'))),
                ('Forks', str(repo_info.get('forks', 'N/A')))
            ]
            
            for metric, value in stats_data:
                row_cells = table.add_row().cells
                row_cells[0].text = metric
                row_cells[1].text = value
        
        # Overview section
        if documentation.get('overview'):
            document.add_heading('Overview', level=1)
            document.add_paragraph(documentation['overview'])
        
        # Architecture section
        if documentation.get('architecture'):
            document.add_heading('High-Level Architecture', level=1)
            document.add_paragraph(documentation['architecture'])
            
            # Architecture diagram
            if documentation.get('architecture_diagram'):
                document.add_heading('Architecture Diagram', level=2)
                document.add_paragraph('Mermaid Diagram Code:')
                p = document.add_paragraph(documentation['architecture_diagram'])
                p.style = 'Intense Quote'
        
        # Modules section
        if documentation.get('modules'):
            document.add_heading('Module Breakdown', level=1)
            document.add_paragraph(documentation['modules'])
        
        # Data Flow section
        if documentation.get('data_flow'):
            document.add_heading('Data Flow Analysis', level=1)
            document.add_paragraph(documentation['data_flow'])
            
            # Data flow diagram
            if documentation.get('data_flow_diagram'):
                document.add_heading('Data Flow Diagram', level=2)
                document.add_paragraph('Mermaid Diagram Code:')
                p = document.add_paragraph(documentation['data_flow_diagram'])
                p.style = 'Intense Quote'
        
        # Design Patterns section
        if documentation.get('design_patterns'):
            document.add_heading('Design Patterns', level=1)
            document.add_paragraph(documentation['design_patterns'])
        
        # Save to buffer
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        
        return buffer
        
    except Exception as e:
        logging.error(f"Error generating DOCX: {e}")
        raise